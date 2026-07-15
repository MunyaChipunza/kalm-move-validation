from __future__ import annotations

import csv
import hashlib
import json
import textwrap
from datetime import datetime, timezone
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


REPO = Path(__file__).resolve().parents[1]
BASE = REPO / "reports/KS-ACTIVE-ARCHIVE/FINAL-PRODUCTION-RELEASE-20260714"
PACK = BASE / "SINGLE-SELLER-STOCK-TRANSFER-20260715"
JAN_ORDER = Path(r"G:\My Drive\Master Folder\08 Business Documents\KS active\January 2023 Order.xlsx")
NOW = datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")

PRODUCT_MANIFEST = BASE / "FINAL-PRODUCT-MANIFEST.json"
SKU_MANIFEST = BASE / "FINAL-SKU-MANIFEST.json"
INVENTORY_MANIFEST = BASE / "FINAL-INVENTORY-MANIFEST.json"
PRICE_MANIFEST = BASE / "FINAL-PRICE-MANIFEST.json"

EXPECTED_HASHES = {
    PRODUCT_MANIFEST.name: "28ED7F0EA89903DF44ED51A5E1144356EABE16304F1B97ED375E6C000E2F91BB",
    SKU_MANIFEST.name: "58EDB97262171552D9ACE68EC67113B836A283B9E2FC85E5B2484F8C38656085",
    INVENTORY_MANIFEST.name: "4CAAC3BB544407718A90BAD56860D4DB85D7BCFBDE355E94B4295292D46E7DB2",
    PRICE_MANIFEST.name: "88DFBD5269191C683E4139988A0C6ED06FD9F39575FF141213B104CE2637DF0B",
}

FALLBACK_COST = 149.35

# Direct family matches to the January 2023 cost workbook. Products with no reliable
# family match deliberately use the user-authorised weighted historical unit cost.
COST_MATCHES = {
    "P002": ("hallow-back-romper", "Direct January 2023 family match"),
    "P003": ("seamless-high-stretch-scrunch-butt-leggings", "Direct January 2023 family match"),
    "P010": ("seamless-crisscross-cut-out-back-sports-bra", "Direct January 2023 family match"),
    "P012": ("seamless-breathable-scrunch-butt-shorts", "Direct January 2023 family match"),
    "P019": ("high-support-seamless-cut-out-sports-bra", "Direct January 2023 family match"),
    "P020": ("mid-support-seamless-crisscross-sports-bra", "Direct January 2023 family match"),
    "P026": ("high-waist-seamless-shorts", "Direct January 2023 family match"),
    "P028": ("high-waist-seamless-leggings", "Direct January 2023 family match"),
    "P030": ("seamless-crisscross-back-sports-bra", "Direct January 2023 family match"),
    "P035": ("seamless-breathable-scrunch-butt-leggings", "Direct January 2023 family match"),
}

PROVISIONAL_CODES = {"P027", "P033", "P049", "P050"}


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest().upper()


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def verify_protected_inputs() -> dict:
    results = {}
    for path in [PRODUCT_MANIFEST, SKU_MANIFEST, INVENTORY_MANIFEST, PRICE_MANIFEST]:
        actual = sha256(path)
        expected = EXPECTED_HASHES[path.name]
        results[path.name] = {
            "path": str(path),
            "expectedSha256": expected,
            "actualSha256": actual,
            "matches": actual == expected,
        }
        if actual != expected:
            raise SystemExit(f"Manifest hash mismatch for {path}: {actual} != {expected}")
    return results


def load_january_costs() -> tuple[dict, list]:
    wb = openpyxl.load_workbook(JAN_ORDER, data_only=True, read_only=True)
    ws = wb["products_export"]
    headers = [c for c in next(ws.iter_rows(min_row=1, max_row=1, values_only=True))]
    costs = {}
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        rec = dict(zip(headers, row))
        title = rec.get("KS Active Title")
        if not title:
            continue
        cost = rec.get("Cost per item")
        if cost is not None:
            costs[str(title)] = float(cost)
        rows.append(rec)
    return costs, rows


def build_transfer_lines() -> tuple[list[dict], dict]:
    sku_manifest = load_json(SKU_MANIFEST)
    products = {p["productCode"]: p for p in load_json(PRODUCT_MANIFEST)["products"]}
    prices = {p["productCode"]: p["price"] for p in load_json(PRICE_MANIFEST)["prices"]}
    jan_costs, jan_rows = load_january_costs()

    lines = []
    for v in sku_manifest["variants"]:
        code = v["productCode"]
        qty = int(v["quantity"])
        if code in COST_MATCHES:
            family, basis_note = COST_MATCHES[code]
            unit_cost = float(jan_costs[family])
            cost_basis = f"January 2023 Order.xlsx family: {family}"
            provisional = False
        else:
            if code not in PROVISIONAL_CODES:
                raise SystemExit(f"No cost basis configured for {code}")
            family = None
            basis_note = "Weighted historical unit cost management estimate"
            unit_cost = FALLBACK_COST
            cost_basis = "Weighted historical unit cost management estimate"
            provisional = True
        line = {
            "productCode": code,
            "productName": v["productName"],
            "productSlug": v["productSlug"],
            "sku": v["sku"],
            "colour": v["colour"],
            "size": v["size"],
            "quantity": qty,
            "unitCostZar": round(unit_cost, 2),
            "lineTransferValueZar": round(unit_cost * qty, 2),
            "approvedSellingPriceZar": float(prices[code]),
            "costBasis": cost_basis,
            "sourceFamily": family or "",
            "basisNote": basis_note,
            "provisional": provisional,
            "seller": "K SENTWA T/A KS ACTIVE",
            "buyer": "KALM COLLECTIVE (PTY) LTD",
            "legalSellerAfterTransfer": "KALM COLLECTIVE (PTY) LTD",
            "brand": "KS Active",
            "category": "KS Active Archive",
            "enabled": bool(v["enabled"]),
            "availability": v["availability"],
        }
        lines.append(line)

    totals = {
        "productCount": len(products),
        "skuCount": len(lines),
        "physicalUnits": sum(l["quantity"] for l in lines),
        "stockedColourCount": sum(len(p["colours"]) for p in products.values()),
        "approvedGeneratedAssets": 224,
        "transferValueZar": round(sum(l["lineTransferValueZar"] for l in lines), 2),
        "directCostSkuCount": sum(1 for l in lines if not l["provisional"]),
        "provisionalCostSkuCount": sum(1 for l in lines if l["provisional"]),
        "januaryOrderPath": str(JAN_ORDER),
        "januaryOrderSha256": sha256(JAN_ORDER),
        "januaryOrderRowsReviewed": len(jan_rows),
        "createdAt": NOW,
    }
    return lines, totals


def write_csv(path: Path, rows: list[dict]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def style_header(ws):
    for cell in ws[1]:
        cell.font = openpyxl.styles.Font(bold=True, color="FFFFFF")
        cell.fill = openpyxl.styles.PatternFill("solid", fgColor="1F4D78")
        cell.alignment = openpyxl.styles.Alignment(wrap_text=True, vertical="center")
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions


def autosize(ws, max_width=48):
    for col in ws.columns:
        letter = col[0].column_letter
        width = min(max_width, max(10, max(len(str(c.value or "")) for c in col) + 2))
        ws.column_dimensions[letter].width = width


def write_xlsx(path: Path, transfer_lines: list[dict], totals: dict) -> None:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Transfer Schedule"
    headers = list(transfer_lines[0].keys())
    ws.append(headers)
    for row in transfer_lines:
        ws.append([row[h] for h in headers])
    style_header(ws)
    for row in ws.iter_rows(min_row=2):
        row[6].number_format = "0"
        row[7].number_format = 'R #,##0.00'
        row[8].number_format = 'R #,##0.00'
        row[9].number_format = 'R #,##0.00'
    autosize(ws)

    s = wb.create_sheet("Summary")
    summary_rows = [
        ("Created at", totals["createdAt"]),
        ("Product count", totals["productCount"]),
        ("Stocked colour count", totals["stockedColourCount"]),
        ("SKU count", totals["skuCount"]),
        ("Physical units", totals["physicalUnits"]),
        ("Approved generated assets", totals["approvedGeneratedAssets"]),
        ("Transfer value ZAR", totals["transferValueZar"]),
        ("Direct cost SKU count", totals["directCostSkuCount"]),
        ("Provisional cost SKU count", totals["provisionalCostSkuCount"]),
        ("January workbook", totals["januaryOrderPath"]),
        ("January workbook SHA-256", totals["januaryOrderSha256"]),
    ]
    s.append(["Metric", "Value"])
    for row in summary_rows:
        s.append(list(row))
    style_header(s)
    autosize(s, max_width=80)

    z = wb.create_sheet("Zoho Items Import")
    zheaders = [
        "Item Name", "SKU", "Product Code", "Brand", "Category", "Colour", "Size",
        "Sales Rate", "Purchase Rate", "Opening Stock", "Inventory Account",
        "COGS Account", "Sales Account", "Vendor", "Description"
    ]
    z.append(zheaders)
    for l in transfer_lines:
        z.append([
            f"{l['productName']} - {l['colour']} - {l['size']}",
            l["sku"], l["productCode"], "KS Active", "KS Active Archive",
            l["colour"], l["size"], l["approvedSellingPriceZar"], l["unitCostZar"],
            l["quantity"], "Inventory Asset", "Cost of Goods Sold", "Product Sales",
            "K SENTWA T/A KS ACTIVE",
            f"KS Active Archive item sold by KALM Collective. Route: /product/{l['productSlug']}",
        ])
    style_header(z)
    autosize(z)

    c = wb.create_sheet("Posting Control")
    c.append(["Control", "Value"])
    controls = [
        ("Preferred posting", "Items + stock/purchase receipt + vendor bill where Zoho supports it"),
        ("Alternative posting", "Opening stock import + management-approved journal"),
        ("Debit", "Inventory Asset"),
        ("Credit", "Related Party Payable - KS Active Stock Transfer"),
        ("Transfer value ZAR", totals["transferValueZar"]),
        ("Do not duplicate", "Do not both import opening stock and post a second inventory debit"),
        ("VAT treatment", "No VAT unless verified VAT registration is confirmed"),
    ]
    for row in controls:
        c.append(list(row))
    style_header(c)
    autosize(c, max_width=72)

    wb.save(path)


def write_zoho_csvs(transfer_lines: list[dict]) -> None:
    item_rows = []
    stock_rows = []
    for l in transfer_lines:
        item_rows.append({
            "Item Name": f"{l['productName']} - {l['colour']} - {l['size']}",
            "SKU": l["sku"],
            "Item Description": f"KS Active Archive item sold by KALM Collective. Product code: {l['productCode']}. Colour: {l['colour']}. Size: {l['size']}. Route: /product/{l['productSlug']}",
            "Selling Price": f"{l['approvedSellingPriceZar']:.2f}",
            "Product Type": "goods",
            "Item Type": "Inventory",
            "Track Inventory": "true",
            "Status": "Active",
            "Unit": "pcs",
            "Sellable": "true",
            "Purchasable": "true",
            "Sales Account": "Product Sales",
            "Purchase Description": f"KS Active Archive stock transfer from K SENTWA T/A KS ACTIVE. Product code: {l['productCode']}.",
            "Purchase Price": f"{l['unitCostZar']:.2f}",
            "Purchase Account": "Cost of Goods Sold",
            "Inventory Account": "Inventory Asset",
            "Opening Stock": l["quantity"],
            "Opening Stock Value": f"{l['lineTransferValueZar']:.2f}",
            "Preferred Vendor": "K SENTWA T/A KS ACTIVE",
            "Brand": "KS Active",
            "CF.Product Code": l["productCode"],
            "CF.Category": "KS Active Archive",
            "CF.Colour": l["colour"],
            "CF.Size": l["size"],
        })
        stock_rows.append({
            "SKU": l["sku"],
            "Item Name": f"{l['productName']} - {l['colour']} - {l['size']}",
            "Opening Stock": l["quantity"],
            "Opening Stock Value": f"{l['lineTransferValueZar']:.2f}",
            "Purchase Rate": f"{l['unitCostZar']:.2f}",
            "Vendor": "K SENTWA T/A KS ACTIVE",
            "Brand": "KS Active",
            "Category": "KS Active Archive",
        })
    write_csv(PACK / "ZOHO_ITEMS_IMPORT.csv", item_rows)
    write_csv(PACK / "ZOHO_OPENING_STOCK_IMPORT.csv", stock_rows)


def add_doc_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(31, 77, 120)
    return p


def add_para(doc, text="", bold_label=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_label:
        r = p.add_run(bold_label)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_signature_block(doc, party):
    add_para(doc, f"Signed for: {party}")
    for label in ["Name", "Capacity", "Signature", "Date", "Place"]:
        add_para(doc, f"{label}: " + "_" * 55)


def save_docx(doc, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def make_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 77, 120)
    p2 = doc.add_paragraph(subtitle)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].italic = True
    add_para(doc, f"Prepared: {NOW}")
    return doc


def write_pdf(path: Path, title: str, sections: list[tuple[str, list[str]]], wide_table=None):
    doc = SimpleDocTemplate(str(path), pagesize=landscape(LETTER) if wide_table else LETTER,
                            rightMargin=0.55*inch, leftMargin=0.55*inch,
                            topMargin=0.55*inch, bottomMargin=0.55*inch)
    styles = getSampleStyleSheet()
    h = ParagraphStyle("KalmH", parent=styles["Heading1"], textColor=colors.HexColor("#1F4D78"))
    h2 = ParagraphStyle("KalmH2", parent=styles["Heading2"], textColor=colors.HexColor("#2E74B5"))
    body = ParagraphStyle("KalmBody", parent=styles["BodyText"], leading=13, spaceAfter=6)
    story = [Paragraph(title, h), Paragraph(f"Prepared: {NOW}", body), Spacer(1, 8)]
    for heading, paras in sections:
        story.append(Paragraph(heading, h2))
        for para in paras:
            story.append(Paragraph(para.replace("&", "&amp;"), body))
        story.append(Spacer(1, 6))
    if wide_table:
        story.append(PageBreak())
        tbl = Table(wide_table["rows"], repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4D78")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(tbl)
    doc.build(story)


def write_signature_docs(transfer_lines: list[dict], totals: dict) -> dict:
    outputs = {}

    agreement = make_doc("KS ACTIVE ARCHIVE STOCK SALE AND TRANSFER AGREEMENT",
                         "Deferred sale and transfer of remaining KS Active Archive inventory")
    add_doc_heading(agreement, "Parties")
    add_para(agreement, "Seller: K SENTWA T/A KS ACTIVE. Registration/tax/address details: [TO BE COMPLETED FROM VERIFIED SOURCE RECORDS].")
    add_para(agreement, "Buyer: KALM COLLECTIVE (PTY) LTD. Registration/tax/address details: [TO BE COMPLETED FROM VERIFIED SOURCE RECORDS].")
    add_doc_heading(agreement, "Commercial Terms")
    terms = [
        f"Transferred inventory: 104 unique SKUs and 111 saleable physical units listed in the transfer schedule.",
        f"Transfer price: R {totals['transferValueZar']:,.2f}, calculated on cost basis and not retail selling price.",
        "Payment treatment: full purchase price remains payable by KALM Collective to K SENTWA T/A KS ACTIVE.",
        "Payment terms: payable on demand or under later written repayment arrangement. No interest unless agreed later in writing.",
        "Title and risk pass to KALM Collective on signature approval and acknowledgement of physical control.",
        "KALM Collective becomes the sole customer-facing seller, merchant of record, customer service party, returns party and refund party.",
        "KS Active remains the product brand. The legal seller on customer invoices is KALM Collective.",
        "No VAT is charged unless verified source records confirm VAT registration and the correct VAT treatment.",
        "Governing law: South Africa. Variations must be in writing.",
    ]
    for t in terms:
        add_para(agreement, t)
    add_doc_heading(agreement, "Limited Warranties")
    add_para(agreement, "The seller warrants only that it is authorised to approve the stock transfer and brand use for the Archive sale, subject to evidence and signature. No unsupported ownership, tax, identity, bank or condition representations are inserted by Codex.")
    add_doc_heading(agreement, "Signatures")
    add_signature_block(agreement, "KALM COLLECTIVE (PTY) LTD")
    add_signature_block(agreement, "K SENTWA T/A KS ACTIVE")
    outputs["agreement_docx"] = PACK / "KS_ACTIVE_ARCHIVE_STOCK_SALE_AND_TRANSFER_AGREEMENT.docx"
    save_docx(agreement, outputs["agreement_docx"])

    resolution = make_doc("KALM BOARD RESOLUTION ACCEPTING KS ACTIVE STOCK",
                          "Management-approved treatment pending signatures")
    add_para(resolution, "Resolved that KALM Collective accepts transfer of the KS Active Archive inventory under the stock sale and transfer agreement, subject to signature by the authorised KS Active representative.")
    add_para(resolution, f"Approved stock scope: 14 products, 56 stocked colours, 104 SKUs, 111 units. Transfer value: R {totals['transferValueZar']:,.2f}.")
    add_para(resolution, "Professional accounting/legal verification may be obtained later but is not a release blocker under Munya's management instruction.")
    add_signature_block(resolution, "KALM COLLECTIVE (PTY) LTD")
    outputs["board_docx"] = PACK / "KALM_BOARD_RESOLUTION_ACCEPTING_KS_ACTIVE_STOCK.docx"
    save_docx(resolution, outputs["board_docx"])

    seller = make_doc("KS ACTIVE SELLER DECLARATION AND TRANSFER AUTHORITY",
                      "For K SENTWA T/A KS ACTIVE approval")
    add_para(seller, "I confirm that I am the owner or authorised representative of K SENTWA T/A KS ACTIVE for the purpose of approving the KS Active Archive stock transfer.")
    add_para(seller, "I authorise the transfer of the inventory listed in the attached schedule to KALM Collective under deferred-payment terms.")
    add_para(seller, "I authorise KALM Collective to sell the transferred stock as legal seller while using KS Active as the product brand for the Archive sale.")
    add_para(seller, "Identity number, address and signature fields are intentionally left fillable and must not be fabricated by Codex.")
    add_signature_block(seller, "K SENTWA T/A KS ACTIVE")
    outputs["seller_docx"] = PACK / "KS_ACTIVE_SELLER_DECLARATION_AND_TRANSFER_AUTHORITY.docx"
    save_docx(seller, outputs["seller_docx"])

    brand = make_doc("KS ACTIVE BRAND LICENCE FOR ARCHIVE SALE",
                     "Limited brand-use permission for the approved Archive sale")
    for text in [
        "KS Active grants KALM Collective a limited licence to use the KS Active name, approved product names, approved product descriptions, verified logo assets and approved generated ecommerce imagery solely for the Archive sale.",
        "This licence does not grant rights beyond the Archive sale unless separately agreed in writing.",
        "KALM Collective must not display KS Active as a separate legal merchant after the stock transfer; KS Active is a product brand.",
        "Private source photographs, Kuhle's likeness, private filenames, signatures and identity details must not be published.",
    ]:
        add_para(brand, text)
    add_signature_block(brand, "K SENTWA T/A KS ACTIVE")
    add_signature_block(brand, "KALM COLLECTIVE (PTY) LTD")
    outputs["brand_docx"] = PACK / "KS_ACTIVE_BRAND_LICENCE_FOR_ARCHIVE_SALE.docx"
    save_docx(brand, outputs["brand_docx"])

    memo = make_doc("KALM ACCOUNTING TREATMENT MEMO",
                    "Deferred stock transfer into single KALM Zoho Books and Inventory organisation")
    for text in [
        "Accounting treatment: debit Inventory Asset and credit Related Party Payable - KS Active Stock Transfer for the approved transfer value.",
        "Do not treat the transfer as sales revenue. Do not post VAT unless verified VAT registration and tax treatment are confirmed.",
        "Zoho Inventory is the operational inventory master after transfer. The storefront and intranet must mirror accepted Zoho export state.",
        "If Zoho supports purchase receipt/vendor bill workflow, use that. If not, use opening stock import plus a management-approved journal without duplicating inventory value.",
        f"Transfer value to post: R {totals['transferValueZar']:,.2f}.",
    ]:
        add_para(memo, text)
    outputs["accounting_docx"] = PACK / "KALM_ACCOUNTING_TREATMENT_MEMO.docx"
    save_docx(memo, outputs["accounting_docx"])

    # Direct PDFs with the same substantive content, not signed.
    common_sections = [
        ("Status", ["Prepared signature-ready draft. Not signed. Kuhle/K SENTWA approval is required before posting or production release."]),
        ("Scope", [f"14 products, 56 stocked colours, 104 SKUs, 111 units, transfer value R {totals['transferValueZar']:,.2f}."]),
    ]
    pdf_specs = [
        ("KS_ACTIVE_ARCHIVE_STOCK_SALE_AND_TRANSFER_AGREEMENT.pdf", "KS Active Archive Stock Sale and Transfer Agreement", [
            ("Parties", ["Seller: K SENTWA T/A KS ACTIVE. Buyer: KALM COLLECTIVE (PTY) LTD. Unverified legal/tax/address fields remain fillable."]),
            ("Terms", terms),
            ("Signatures", ["Unsigned. Signature blocks are in the editable DOCX."]),
        ]),
        ("KALM_BOARD_RESOLUTION_ACCEPTING_KS_ACTIVE_STOCK.pdf", "KALM Board Resolution Accepting KS Active Stock", common_sections),
        ("KS_ACTIVE_SELLER_DECLARATION_AND_TRANSFER_AUTHORITY.pdf", "KS Active Seller Declaration and Transfer Authority", common_sections),
        ("KS_ACTIVE_BRAND_LICENCE_FOR_ARCHIVE_SALE.pdf", "KS Active Brand Licence for Archive Sale", common_sections),
        ("KALM_ACCOUNTING_TREATMENT_MEMO.pdf", "KALM Accounting Treatment Memo", common_sections),
        ("ZOHO_POSTING_CONTROL_MEMO.pdf", "Zoho Posting Control Memo", common_sections),
        ("MANAGEMENT_APPROVAL_RECORD.pdf", "Management Approval Record", [
            ("KALM-side authority", ["Munya authorised KALM-side treatment and release path in the controlling prompt."]),
            ("Outstanding approval", ["Kuhle/K SENTWA approval must be captured before transferDocumentsApproved or stockOwnershipTransferredToKALM can be true."]),
        ]),
    ]
    for filename, title, sections in pdf_specs:
        path = PACK / filename
        write_pdf(path, title, sections)
        outputs[filename] = path

    schedule_rows = [["Product", "SKU", "Colour", "Size", "Qty", "Unit cost", "Line value"]]
    for l in transfer_lines:
        schedule_rows.append([l["productCode"], l["sku"], l["colour"], l["size"], str(l["quantity"]), f"R {l['unitCostZar']:.2f}", f"R {l['lineTransferValueZar']:.2f}"])
    write_pdf(PACK / "KS_ACTIVE_ARCHIVE_INVENTORY_TRANSFER_SCHEDULE.pdf",
              "KS Active Archive Inventory Transfer Schedule",
              [("Schedule note", ["The complete 104-SKU schedule is included below. The editable schedule is also provided in XLSX/CSV format."])],
              wide_table={"rows": schedule_rows})
    outputs["schedule_pdf"] = PACK / "KS_ACTIVE_ARCHIVE_INVENTORY_TRANSFER_SCHEDULE.pdf"
    return outputs


def write_markdown_and_json(transfer_lines, totals, hashes, outputs):
    method = f"""# Transfer Valuation Method

Prepared: {NOW}

## Source hierarchy

1. Final physical inventory manifest: `{INVENTORY_MANIFEST}`
2. Final SKU manifest: `{SKU_MANIFEST}`
3. January 2023 Order workbook: `{JAN_ORDER}`
4. User-authorised weighted historical unit cost estimate for products without reliable direct family match.

## Cost method

Transfer value uses historical cost, not retail launch price. Direct product-family matches use the January 2023 workbook `Cost per item`. P027, P033, P049 and P050 use the authorised weighted historical unit cost estimate of R {FALLBACK_COST:.2f} because no reliable direct family match was identified in the January workbook during this pass.

## Result

- Product count: {totals['productCount']}
- Stocked colours: {totals['stockedColourCount']}
- SKU count: {totals['skuCount']}
- Physical units: {totals['physicalUnits']}
- Transfer value: R {totals['transferValueZar']:,.2f}
- Direct-cost SKUs: {totals['directCostSkuCount']}
- Provisional-cost SKUs: {totals['provisionalCostSkuCount']}

## Release control

The transfer pack is draft/signature-ready only. Kuhle/K SENTWA approval must be captured before any Zoho stock posting, ownership transfer, or production release is marked complete.
"""
    (PACK / "TRANSFER-VALUATION-METHOD.md").write_text(method, encoding="utf-8")

    validation = {
        "createdAt": NOW,
        "passed": totals["skuCount"] == 104 and totals["physicalUnits"] == 111 and len({l["sku"] for l in transfer_lines}) == 104 and totals["transferValueZar"] > 0,
        "checks": {
            "manifestHashesMatch": all(v["matches"] for v in hashes.values()),
            "uniqueSkuCount": len({l["sku"] for l in transfer_lines}),
            "expectedSkuCount": 104,
            "physicalUnits": totals["physicalUnits"],
            "expectedPhysicalUnits": 111,
            "noNegativeQuantity": all(l["quantity"] >= 0 for l in transfer_lines),
            "everyUnitHasCostBasis": all(l["unitCostZar"] > 0 and l["costBasis"] for l in transfer_lines),
            "transferValueEqualsLineSum": round(sum(l["lineTransferValueZar"] for l in transfer_lines), 2) == totals["transferValueZar"],
        },
        "transferValueZar": totals["transferValueZar"],
        "provisionalCostProductCodes": sorted(PROVISIONAL_CODES),
        "blockingGate": "Kuhle/K SENTWA signed approval is required before transferDocumentsApproved=true and stockOwnershipTransferredToKALM=true.",
    }
    (PACK / "TRANSFER-VALUATION-VALIDATION.json").write_text(json.dumps(validation, indent=2), encoding="utf-8")

    control = {
        "createdAt": NOW,
        "operatingModel": {
            "legalSeller": "KALM COLLECTIVE (PTY) LTD",
            "merchantOfRecord": "KALM COLLECTIVE (PTY) LTD",
            "inventoryOwnerAfterTransfer": "KALM COLLECTIVE (PTY) LTD",
            "brand": "KS Active",
            "accountingOrganisation": "KALM Collective Zoho Books organisation",
            "inventoryOrganisation": "KALM Collective Zoho Inventory organisation",
        },
        "sourceEvidence": {
            "manifestHashes": hashes,
            "januaryOrderPath": str(JAN_ORDER),
            "januaryOrderSha256": totals["januaryOrderSha256"],
            "zohoOrganisationObserved": "KALM Collective",
            "zohoOrganisationIdObserved": "930770020",
            "unverifiedFieldsLeftBlank": [
                "company registration numbers",
                "tax numbers",
                "identity numbers",
                "registered addresses",
                "bank details",
                "directors/shareholders/legal representatives",
                "VAT registration status",
                "signatures",
            ],
        },
        "totals": totals,
        "documents": {k: str(v) for k, v in outputs.items()},
        "status": {
            "transferDocumentsApproved": False,
            "stockOwnershipTransferredToKALM": False,
            "zohoItemsCreated": False,
            "zohoBooksPostingComplete": False,
            "intranetReconciled": False,
            "storefrontReconciled": False,
            "releaseAuthorised": False,
            "blockingGates": ["Kuhle/K SENTWA approval not captured"],
        },
    }
    (PACK / "TRANSFER-PACK-CONTROL.json").write_text(json.dumps(control, indent=2), encoding="utf-8")


def main():
    PACK.mkdir(parents=True, exist_ok=True)
    hashes = verify_protected_inputs()
    transfer_lines, totals = build_transfer_lines()

    write_csv(PACK / "KS-ACTIVE-ARCHIVE-TRANSFER-VALUATION.csv", transfer_lines)
    write_zoho_csvs(transfer_lines)
    write_xlsx(PACK / "KS-ACTIVE-ARCHIVE-TRANSFER-VALUATION.xlsx", transfer_lines, totals)
    write_xlsx(PACK / "KS_ACTIVE_ARCHIVE_INVENTORY_TRANSFER_SCHEDULE.xlsx", transfer_lines, totals)
    write_xlsx(PACK / "ZOHO_POSTING_AND_IMPORT_PACK.xlsx", transfer_lines, totals)
    outputs = write_signature_docs(transfer_lines, totals)
    write_markdown_and_json(transfer_lines, totals, hashes, outputs)

    manifest = {
        "createdAt": NOW,
        "packPath": str(PACK),
        "transferValueZar": totals["transferValueZar"],
        "files": [],
    }
    for path in sorted(PACK.glob("*")):
        if path.is_file():
            manifest["files"].append({"path": str(path), "sha256": sha256(path), "bytes": path.stat().st_size})
    (PACK / "PACK-FILE-MANIFEST.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(json.dumps({"pack": str(PACK), "transferValueZar": totals["transferValueZar"], "files": len(manifest["files"])}, indent=2))


if __name__ == "__main__":
    main()
